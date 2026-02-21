"""
Resume Service — parse resumes from PDF/DOCX uploads.

Responsibilities:
  • Extract raw text from PDF (pdfplumber) or DOCX (python-docx)
  • Clean up extracted text
  • Send to LLM for structured section extraction
  • Cache parsed results in memory (MVP)
"""

from __future__ import annotations

import logging
import uuid
from typing import Any

import pdfplumber
import docx
from io import BytesIO

from app.models.resume_models import (
    ParsedResume,
    ContactInfo,
    ExperienceEntry,
    ProjectEntry,
    EducationEntry,
)
from app.prompts.resume_extractor import SYSTEM_PROMPT, USER_PROMPT_TEMPLATE
from app.services.llm_service import complete_json
from app.utils.file_hash import md5_hash
from app.utils.text_cleanup import normalize_text
from app.services import vector_store

logger = logging.getLogger(__name__)

# In-memory resume cache (session-scoped, lost on restart — fine for MVP)
_resume_cache: dict[str, ParsedResume] = {}


# ── Public API ───────────────────────────────────────────────────────────────


async def parse_resume(
    *,
    file_bytes: bytes,
    file_name: str,
    provider: str,
    model_key: str,
    api_key: str,
    source: str = "local_upload",
    drive_file_id: str | None = None,
) -> ParsedResume:
    """Parse an uploaded resume file (PDF/DOCX) into structured ParsedResume."""
    file_hash = md5_hash(file_bytes)

    # Check cache by hash — avoid reprocessing the same file
    for cached in _resume_cache.values():
        if cached.file_hash == file_hash:
            logger.info(f"Cache hit for {file_name} (hash={file_hash[:8]})")
            return cached

    # Extract raw text
    ext = file_name.rsplit(".", 1)[-1].lower() if "." in file_name else ""
    if ext == "pdf":
        raw_text = _extract_pdf_text(file_bytes)
    elif ext in ("docx", "doc"):
        raw_text = _extract_docx_text(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: .{ext}. Please upload PDF or DOCX.")

    if not raw_text or len(raw_text.strip()) < 50:
        raise ValueError(
            f"Could not extract meaningful text from {file_name}. "
            "The file may be image-based or corrupted."
        )

    # Clean up
    raw_text = normalize_text(raw_text)

    # LLM extraction
    messages = [
        {"role": "system", "content": SYSTEM_PROMPT},
        {"role": "user", "content": USER_PROMPT_TEMPLATE.format(resume_text=raw_text)},
    ]

    logger.info(f"Parsing resume '{file_name}' ({len(raw_text)} chars) with {provider}/{model_key}")

    data = await complete_json(
        provider=provider,
        model_key=model_key,
        api_key=api_key,
        messages=messages,
        prompt_name="resume_extractor",
    )

    parsed = _build_parsed_resume(
        data=data,
        raw_text=raw_text,
        file_name=file_name,
        file_hash=file_hash,
        source=source,
        drive_file_id=drive_file_id,
    )

    _resume_cache[parsed.id] = parsed
    logger.info(f"Parsed resume: id={parsed.id} name={parsed.name} skills={len(parsed.skills)}")

    # Store embedding in vector store for semantic search
    try:
        sections = {}
        if parsed.summary:
            sections["summary"] = parsed.summary
        if parsed.skills:
            sections["skills"] = ", ".join(parsed.skills)
        exp_text = " ".join(
            f"{e.title} at {e.company}. " + " ".join(e.bullets)
            for e in parsed.experience
        )
        if exp_text.strip():
            sections["experience"] = exp_text

        vector_store.store_resume(
            resume_id=parsed.id,
            file_name=file_name,
            file_hash=file_hash,
            source=source,
            full_text=raw_text,
            sections=sections,
        )
    except Exception as e:
        logger.warning(f"Failed to store resume in vector store (non-fatal): {e}")

    return parsed


def get_cached_resume(resume_id: str) -> ParsedResume | None:
    """Retrieve a previously parsed resume by ID."""
    return _resume_cache.get(resume_id)


def list_cached_resumes() -> list[ParsedResume]:
    """Return all cached parsed resumes."""
    return list(_resume_cache.values())


def clear_cached_resumes() -> int:
    """Clear all cached resumes. Returns count cleared."""
    count = len(_resume_cache)
    _resume_cache.clear()
    return count


# ── Text Extraction ──────────────────────────────────────────────────────────


def _extract_pdf_text(file_bytes: bytes) -> str:
    """Extract text from a PDF file using pdfplumber."""
    text_parts: list[str] = []
    with pdfplumber.open(BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n\n".join(text_parts)


def _extract_docx_text(file_bytes: bytes) -> str:
    """Extract text from a DOCX file using python-docx."""
    document = docx.Document(BytesIO(file_bytes))
    text_parts: list[str] = []

    for para in document.paragraphs:
        stripped = para.text.strip()
        if stripped:
            text_parts.append(stripped)

    # Also extract text from tables (some resumes use tables for layout)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    text_parts.append(cell_text)

    return "\n".join(text_parts)


# ── Helpers ──────────────────────────────────────────────────────────────────


def _build_parsed_resume(
    data: dict | list,
    raw_text: str,
    file_name: str,
    file_hash: str,
    source: str = "local_upload",
    drive_file_id: str | None = None,
) -> ParsedResume:
    """Build a ParsedResume from LLM JSON output, with safe defaults."""
    if isinstance(data, list):
        data = data[0] if data else {}

    d: dict[str, Any] = data  # type: ignore

    # Build contact info
    contact_data = d.get("contact", {})
    if not isinstance(contact_data, dict):
        contact_data = {}

    contact = ContactInfo(
        email=contact_data.get("email"),
        phone=contact_data.get("phone"),
        linkedin=contact_data.get("linkedin"),
        location=contact_data.get("location"),
    )

    # Build experience entries
    experience = []
    for exp in _ensure_list_of_dicts(d.get("experience")):
        experience.append(ExperienceEntry(
            title=exp.get("title", "Unknown Role"),
            company=exp.get("company", "Unknown Company"),
            dates=exp.get("dates", ""),
            bullets=_ensure_str_list(exp.get("bullets")),
        ))

    # Build project entries
    projects = []
    for proj in _ensure_list_of_dicts(d.get("projects")):
        projects.append(ProjectEntry(
            name=proj.get("name", "Unnamed Project"),
            technologies=_ensure_str_list(proj.get("technologies")),
            bullets=_ensure_str_list(proj.get("bullets")),
        ))

    # Build education entries
    education = []
    for edu in _ensure_list_of_dicts(d.get("education")):
        education.append(EducationEntry(
            degree=edu.get("degree", "Unknown Degree"),
            school=edu.get("school", "Unknown School"),
            year=edu.get("year"),
        ))

    return ParsedResume(
        id=str(uuid.uuid4()),
        file_name=file_name,
        file_hash=file_hash,
        source=source,
        drive_file_id=drive_file_id,
        name=d.get("name", "Unknown"),
        contact=contact,
        tagline=d.get("tagline"),
        summary=d.get("summary"),
        skills=_ensure_str_list(d.get("skills")),
        experience=experience,
        projects=projects,
        education=education,
        certifications=_ensure_str_list(d.get("certifications")),
        raw_text=raw_text,
    )


def _ensure_str_list(val: Any) -> list[str]:
    """Ensure the value is a list of strings."""
    if isinstance(val, list):
        return [str(v) for v in val]
    if isinstance(val, str):
        return [val]
    return []


def _ensure_list_of_dicts(val: Any) -> list[dict]:
    """Ensure the value is a list of dicts."""
    if isinstance(val, list):
        return [v for v in val if isinstance(v, dict)]
    return []
