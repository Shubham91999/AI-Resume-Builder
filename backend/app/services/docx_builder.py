"""
DOCX Builder Service — generate a professional resume document from TailoredResume.

Uses python-docx to create a clean, ATS-friendly single-column resume with:
  - Name + contact info header
  - Tagline
  - Professional Summary
  - Skills (categorized)
  - Experience (with bullet points)
  - Projects (selected from bank)
  - Education
  - Certifications
"""

from __future__ import annotations

import io
import logging
from typing import Optional

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from app.models.tailor_models import TailoredResume

logger = logging.getLogger(__name__)


def build_docx(tailored: TailoredResume) -> io.BytesIO:
    """
    Generate a DOCX resume from a TailoredResume model.

    Returns a BytesIO buffer containing the DOCX file.
    """
    doc = Document()

    # ── Page margins ─────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    # ── Default font ─────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Reduce paragraph spacing globally
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(2)

    # ── Name ─────────────────────────────────────────────────────────────
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(tailored.name.upper())
    name_run.bold = True
    name_run.font.size = Pt(16)
    name_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    name_para.paragraph_format.space_after = Pt(2)

    # ── Contact Line ─────────────────────────────────────────────────────
    contact_parts = []
    contact = tailored.contact
    if isinstance(contact, dict):
        if contact.get("email"):
            contact_parts.append(contact["email"])
        if contact.get("phone"):
            contact_parts.append(contact["phone"])
        if contact.get("linkedin"):
            contact_parts.append(contact["linkedin"])
        if contact.get("location"):
            contact_parts.append(contact["location"])

    if contact_parts:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_run = contact_para.add_run(" | ".join(contact_parts))
        contact_run.font.size = Pt(9)
        contact_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        contact_para.paragraph_format.space_after = Pt(4)

    # ── Tagline ──────────────────────────────────────────────────────────
    if tailored.tagline:
        tag_para = doc.add_paragraph()
        tag_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tag_run = tag_para.add_run(tailored.tagline)
        tag_run.font.size = Pt(10)
        tag_run.italic = True
        tag_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        tag_para.paragraph_format.space_after = Pt(6)

    # ── Section: Summary ─────────────────────────────────────────────────
    if tailored.summary:
        _add_section_heading(doc, "PROFESSIONAL SUMMARY")
        sum_para = doc.add_paragraph(tailored.summary)
        sum_para.paragraph_format.space_after = Pt(4)

    # ── Section: Skills ──────────────────────────────────────────────────
    if tailored.skills:
        _add_section_heading(doc, "TECHNICAL SKILLS")
        for category, skill_str in tailored.skills.items():
            skill_para = doc.add_paragraph()
            cat_run = skill_para.add_run(f"{category}: ")
            cat_run.bold = True
            cat_run.font.size = Pt(10)
            skill_run = skill_para.add_run(skill_str)
            skill_run.font.size = Pt(10)
            skill_para.paragraph_format.space_after = Pt(1)

    # ── Section: Experience ──────────────────────────────────────────────
    if tailored.experience:
        _add_section_heading(doc, "PROFESSIONAL EXPERIENCE")
        for exp in tailored.experience:
            # Title + Company line
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(f"{exp.title}")
            title_run.bold = True
            title_run.font.size = Pt(10)

            sep_run = title_para.add_run(f" — {exp.company}")
            sep_run.font.size = Pt(10)
            sep_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

            # Dates (right-aligned via tab)
            if exp.dates:
                date_run = title_para.add_run(f"\t{exp.dates}")
                date_run.font.size = Pt(9)
                date_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                # Add a right tab stop
                tab_stops = title_para.paragraph_format.tab_stops
                tab_stops.add_tab_stop(Inches(6.3), alignment=2)  # 2 = RIGHT

            title_para.paragraph_format.space_before = Pt(6)
            title_para.paragraph_format.space_after = Pt(1)

            # Bullets
            for bullet in exp.bullets:
                bp = doc.add_paragraph(style="List Bullet")
                bp.text = ""
                br = bp.add_run(bullet)
                br.font.size = Pt(9.5)
                bp.paragraph_format.space_before = Pt(0)
                bp.paragraph_format.space_after = Pt(1)
                bp.paragraph_format.left_indent = Inches(0.25)

    # ── Section: Projects ────────────────────────────────────────────────
    if tailored.projects:
        _add_section_heading(doc, "PROJECTS")
        for proj in tailored.projects:
            proj_para = doc.add_paragraph()
            proj_run = proj_para.add_run(proj.name)
            proj_run.bold = True
            proj_run.font.size = Pt(10)
            proj_para.paragraph_format.space_before = Pt(4)
            proj_para.paragraph_format.space_after = Pt(1)

            for bullet in proj.bullets:
                bp = doc.add_paragraph(style="List Bullet")
                bp.text = ""
                br = bp.add_run(bullet)
                br.font.size = Pt(9.5)
                bp.paragraph_format.space_before = Pt(0)
                bp.paragraph_format.space_after = Pt(1)
                bp.paragraph_format.left_indent = Inches(0.25)

    # ── Section: Education ───────────────────────────────────────────────
    if tailored.education:
        _add_section_heading(doc, "EDUCATION")
        for edu in tailored.education:
            if isinstance(edu, dict):
                degree = edu.get("degree", "")
                school = edu.get("school", "")
                year = edu.get("year", "")
            else:
                degree = getattr(edu, "degree", "")
                school = getattr(edu, "school", "")
                year = getattr(edu, "year", "")

            edu_para = doc.add_paragraph()
            deg_run = edu_para.add_run(degree)
            deg_run.bold = True
            deg_run.font.size = Pt(10)

            if school:
                sch_run = edu_para.add_run(f" — {school}")
                sch_run.font.size = Pt(10)
                sch_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

            if year:
                yr_run = edu_para.add_run(f"  ({year})")
                yr_run.font.size = Pt(9)
                yr_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

            edu_para.paragraph_format.space_after = Pt(2)

    # ── Section: Certifications ──────────────────────────────────────────
    if tailored.certifications:
        _add_section_heading(doc, "CERTIFICATIONS")
        for cert in tailored.certifications:
            cert_para = doc.add_paragraph(f"• {cert}")
            cert_para.paragraph_format.space_after = Pt(1)

    # ── Write to buffer ──────────────────────────────────────────────────
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    logger.info(f"DOCX generated for tailored resume: {tailored.id}")
    return buffer


def _add_section_heading(doc: Document, title: str) -> None:
    """Add a styled section heading with a bottom border."""
    para = doc.add_paragraph()
    run = para.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after = Pt(3)

    # Add bottom border via XML
    pPr = para._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(
        qn("w:bottom"),
        {
            qn("w:val"): "single",
            qn("w:sz"): "4",
            qn("w:space"): "1",
            qn("w:color"): "999999",
        },
    )
    pBdr.append(bottom)
    pPr.append(pBdr)
